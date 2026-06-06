# -*- coding: utf-8 -*-
"""Generate OpenFIRE Gun + Dongle user manual (Chinese, OLED_MENU_ZH)."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = r"d:\code\OpenFIRE-Firmware-ESP32\docs\OpenFIRE光枪与Dongle使用说明书.docx"


def set_cn_font(run, name="微软雅黑", size=11):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_title(doc, text, level=0):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_cn_font(r, size=22)
        r.bold = True
    else:
        doc.add_heading(text, level=level)


def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r)
    r.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            set_cn_font(r)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_cn_font(r, size=10)
                r.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    set_cn_font(r, size=10)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    add_title(doc, "OpenFIRE 光枪与无线 Dongle\n详细使用说明书")
    add_p(doc, "适用固件：OpenFIRE-Firmware-ESP32（ESP32-S3 光枪 + ESP-NOW Dongle，中文 OLED 菜单版）")
    add_p(doc, "配套配置软件：OpenFIRE-App-cn（可选）")
    add_p(doc, "文档版本：1.1  |  编写日期：2026年6月")
    add_p(doc, "说明：本版说明书与固件 OLED_MENU_ZH 中文界面一一对应；表中「OLED 显示」列为屏幕上实际文字。")
    doc.add_paragraph()

    # --- 1 ---
    add_title(doc, "一、产品概述", 1)
    add_p(doc,
          "OpenFIRE 是一套基于四红外点（4IR）定位的光枪系统。ESP32-S3 版本支持两种与 PC 连接方式："
          "① USB 有线直连；② 通过 ESP-NOW 无线 Dongle 连接。无线模式下，PC 识别到的设备与有线连接相同，"
          "游戏与力反馈软件无需额外适配。")
    add_p(doc, "本说明书涵盖：光枪（GUN）日常操作、中文 OLED 暂停菜单、校准、多玩家编号；"
          "无线接收器（Dongle）与无线踏板的配对、屏幕状态与故障排除。", bold=False)

    add_title(doc, "1.1 包装与配件（参考）", 2)
    add_bullets(doc, [
        "光枪本体（含 IR 摄像头、按键、OLED、可选震动/电磁阀等）",
        "USB 数据线（用于有线连接或充电/刷机）",
        "无线 Dongle（ESP32-S3，插 PC USB，SSD1306/TFT 型号显示中文配对状态）",
        "IR 发射器条（需安装在显示器上下或四边，依方形/菱形布局而定）",
        "（可选）无线踏板模块，用于 Time Crisis 类游戏的掩护操作",
    ])

    add_title(doc, "1.2 安全提示", 2)
    add_bullets(doc, [
        "首次开机及摇杆自动校准时，约 2 秒内请勿触碰模拟摇杆。",
        "电磁阀（Solenoid）工作时会有机械冲击与热量，请勿将枪口对准人员或动物。",
        "长时间使用请关注 OLED 温度显示；超温保护触发时会限制力反馈输出。",
        "刷机、改线请在断电状态下进行，并核对板型引脚文档（docs/BOARDS_PIN.md）。",
    ])

    # --- 2 ---
    add_title(doc, "二、首次开机与连接", 1)

    add_title(doc, "2.1 启动顺序（必读）", 2)
    add_p(doc, "光枪上电后按以下顺序自动执行（OLED 中文提示对照）：")
    add_table(doc,
              ["阶段", "OLED 显示（中文）", "说明"],
              [
                  ["摇杆自动校准", "校准就绪", "约 2 秒，请勿动摇杆"],
                  ["USB 检测", "USB模式 → 等待USB", "判断有线或无线"],
                  ["连接成功", "连接就绪 / USB已连接", "有线或无线握手完成"],
                  ["首次屏幕校准", "欢迎! / 扣扳机 / 扣扳机开始校准", "未校准时按扳机进入五点校准"],
              ])

    add_title(doc, "2.2 有线 USB 连接", 2)
    add_bullets(doc, [
        "用 USB 线连接光枪与 PC。",
        "等待 Windows 识别 HID 设备（设备名通常为 FIRECon P1～P4，取决于枪号设置）。",
        "OLED 顶部可能显示 USB已连接；左下角显示 USB 图标。",
        "可用 OpenFIRE-App-cn 连接 COM 口进行高级配置（需进入 Docked 模式，见第十章）。",
    ])

    add_title(doc, "2.3 无线 Dongle 连接（推荐顺序）", 2)
    add_bullets(doc, [
        "步骤 1：将 Dongle 插入 PC USB，等待其完成启动（OLED 显示 WiFi扫描 / 优选信道 / 请稍候，约 10～15 秒）。",
        "步骤 2：光枪不要插 PC USB，然后打开光枪电源。",
        "步骤 3：Dongle 显示 信道xx配对 / 请开光枪 / ESP-NOW监听 时，光枪应开始广播配对。",
        "光枪 OLED 顶部可能显示 信道xx等待。",
        "步骤 4：配对成功后 Dongle 显示 射频已连接 → 玩家 / 信道 / 枪 MAC；PC 出现与光枪配置一致的设备。",
        "步骤 5：若有无线踏板，Dongle 可能短暂显示 等待踏板；光枪搜索踏板时显示 踏台搜索xx秒。",
        "步骤 6：OLED 左下角显示 Wi-Fi 图标表示无线模式。",
        "再次开机时，光枪会优先连接上次配对的 Dongle；若失败则重新全信道搜索。",
    ])

    add_title(doc, "2.4 玩具模式（原 Boy Mode，无 PC 试玩）", 2)
    add_p(doc,
          "无线未连接且 USB 未挂载时，上电约 3 秒内按住扳机可进入玩具模式："
          "OLED 显示 玩具模式?Xs 倒计时，就绪后显示 玩具模式就绪。"
          "可在无 PC 情况下本地试玩电磁阀等功能（具体取决于硬件配置）。")

    # --- 3 OLED ---
    add_title(doc, "三、OLED 状态栏说明（运行界面）", 1)
    add_table(doc,
              ["区域", "OLED 显示示例", "含义"],
              [
                  ["顶部左侧", "P1: P_A 方 42°C", "玩家编号、当前配置简写、布局方/菱、温度"],
                  ["顶部右侧", "关 / 计时5分", "游戏计时：关=未启用；数字=剩余秒/分"],
                  ["顶部整行", "配置: ProfileA P1", "暂停/切换配置时显示"],
                  ["左下", "USB / Wi-Fi 图标", "有线或无线 Dongle"],
                  ["中下", "RF / LOW / AF", "震动FFB、少键模式、连发（缩写，与英文固件相同）"],
                  ["右下", "鼠标 / 手柄 / MiSTer 图标", "当前 USB 输出模式"],
              ])
    add_p(doc, "操作反馈时顶部栏会短暂显示中文，例如：模式:手柄、连发:开、少键:关、震动FFB:开、已发Esc、游戏时间到 等。")

    # --- 4 Pause ---
    add_title(doc, "四、暂停菜单（中文 OLED）", 1)
    add_p(doc,
          "暂停菜单用于切换配置、校准、改 USB 模式、设置枪号、保存设置等。"
          "系统有两种交互方式，由 simplePause 开关决定（默认一般为 Direct Pause 热键模式）。")

    add_title(doc, "4.1 进入与退出暂停", 2)
    add_table(doc,
              ["操作", "按键组合", "OLED 相关提示"],
              [
                  ["进入暂停", "Reload + Select（换弹 + 选择）", "使用 + 当前配置名"],
                  ["退出暂停", "Reload + Home（换弹 + 主页）", "恢复运行界面"],
                  ["长按进入暂停（需开启 Hold to Pause）", "Trigger + A 按住约 2.5 秒", "—"],
                  ["Simple 模式长按退出", "A + B 按住约一半 Hold to Pause 时间", "—"],
              ])

    add_title(doc, "4.2 运行模式快捷键（无需进暂停）", 2)
    add_table(doc,
              ["功能", "按键组合", "OLED 提示"],
              [
                  ["发送 Esc 到 PC", "Reload + Start", "已发Esc"],
                  ["发送 Esc 到 PC（快捷）", "Trigger + A + C（扳机 + A + C 同时松开）", "已发Esc"],
                  ["鼠标模式 ↔ 手柄模式", "A + B + Trigger", "模式:键鼠 / 模式:手柄"],
              ])
    add_p(doc, "注：C 键在固件中对应 Gun C（Reload 按键位）。MiSTer 模式需在暂停菜单「模式切换」中循环选择。")

    add_title(doc, "4.3 Direct Pause（默认热键模式）", 2)
    add_p(doc, "进入暂停后，OLED 四行显示 A> / B> / 始> / 选> 与四个配置名；多数功能通过组合键完成。")
    add_table(doc,
              ["功能", "按键组合"],
              [
                  ["屏幕校准", "暂停内按 Trigger（扳机）"],
                  ["保存全部设置到 Flash", "Start + Select"],
                  ["IR 灵敏度 +", "B + Up"],
                  ["IR 灵敏度 −", "B + Down"],
                  ["运行模式 Normal", "Start + A"],
                  ["运行模式 Average / Average2", "Start + B"],
                  ["发送 Esc", "Reload + Start 或 Trigger+A+C"],
                  ["鼠标 ↔ 手柄", "A + B + Trigger"],
                  ["切换配置 A/B/Start/Select", "暂停界面按对应面键"],
                  ["校准中跳过中心步", "校准过程中按 A"],
              ])

    add_title(doc, "4.4 Simple Pause（列表菜单模式）", 2)
    add_p(doc, "若已在 OpenFIRE App 中开启 Simple Pause：A=上一项，B=下一项，Trigger=确认执行。"
          "下列「OLED 显示」与固件中文菜单完全一致。")
    add_table(doc,
              ["OLED 显示（中文）", "功能说明"],
              [
                  ["中心校准", "以当前位置为摇杆新中心（需配置模拟摇杆引脚）"],
                  ["摇杆校准", "约 4 秒摇杆画圈 + 回中（OLED 提示：摇杆:旋转 → 松开回中 → 摇杆:完成）"],
                  ["屏幕校准", "进入屏幕五点瞄准校准"],
                  ["配置选择", "选择 Profile A/B/Start/Select（子界面显示 A>/B>/始>/选>）"],
                  ["模式切换", "循环：键鼠 → 手柄 → MiSTer（当前:键鼠 / 当前:手柄 / 当前:MiSTer）"],
                  ["布局切换", "方形 ↔ 菱形（布局:方形 / 布局:菱形）"],
                  ["枪号 / 枪号P1-P4", "设置玩家 P1～P4，同步 USB PID 与设备名"],
                  ["摇杆模式", "摇杆输出：摇杆:手柄 / 摇杆:十字 / 摇杆:按键"],
                  ["死区", "摇杆死区 0～30%（死区:N%）"],
                  ["轴模式", "有符号 / 无符号轴"],
                  ["互换摇杆", "红外定位轴与物理摇杆左右互换"],
                  ["游戏计时", "0 / 5 / 10 / 15 / 20 分钟（计时:关 / 计时N分），到期显示 游戏时间到"],
                  ["保存设置", "保存配置与全局设置到 Flash（保存中... / 保存 成功）"],
                  ["震动开关", "电机震动总开关（震动:开 / 震动:关）"],
                  ["电磁开关", "电磁阀总开关（电磁:开 / 电磁:关）"],
                  ["连发开关", "连发模式（连发:开 / 连发:关）"],
                  ["少键模式", "低按键映射模式（少键:开 / 少键:关）"],
                  ["震动FFB开关", "是否用震动代替电磁响应串口力反馈（震动FFB:开 / 关 / 无）"],
                  ["反转X轴 / 反转Y轴", "模拟摇杆轴向反转"],
                  ["发送Esc", "向 PC 发送一次 Esc 键"],
              ])
    add_p(doc, "注：无对应硬件的菜单项会自动跳过。枪号子菜单：A/B 选择 P1～P4，Trigger 保存，Reload+Home 取消。")

    # --- 5 Gun ID ---
    add_title(doc, "五、枪号与多光枪设置（P1～P4）", 1)
    add_bullets(doc, [
        "每只光枪可设为 P1、P2、P3、P4，对应 USB Product ID 1～4，设备名如 FIRECon P1。",
        "OLED 显示 枪号P1～P4 或顶部 枪号P1；Simple 菜单选 枪号；或 OpenFIRE-App-cn → Settings → USB Identity。",
        "修改后需 保存设置，并重新插拔 USB 后 Windows 才会识别新编号。",
        "多枪同屏：每只枪设不同枪号，各接不同 USB 口或使用多个 Dongle。",
    ])

    # --- 6 USB modes ---
    add_title(doc, "六、USB 输出模式", 1)
    add_table(doc,
              ["模式", "OLED 显示", "典型用途"],
              [
                  ["键鼠（默认）", "当前:键鼠 / 模式:键鼠", "MAME、Demul、TeknoParrot 等"],
                  ["手柄", "当前:手柄 / 模式:手柄", "需要手柄协议的游戏"],
                  ["MiSTer", "当前:MiSTer / 模式:MiSTer", "MiSTer FPGA 平台"],
              ])
    add_p(doc, "切换方式：Simple 菜单「模式切换」；运行中 A+B+Trigger（键鼠↔手柄）；OpenFIRE App 设置。")

    # --- 7 Profiles ---
    add_title(doc, "七、校准配置文件（Profile）", 1)
    add_p(doc, "系统内置 4 套 Profile，默认名称：Profile A、Profile B、Profile Start、Profile Select。")
    add_p(doc, "OLED 可能显示为 P_A、P_B 等简写，或 配置:名称；配置选择界面显示 名称(方) / 名称(菱)。")
    add_bullets(doc, [
        "Direct 暂停：按 OLED 上 A>/B>/始>/选> 对应键快速切换。",
        "Simple 暂停：「配置选择」子菜单。",
        "OpenFIRE App：Profiles 标签页编辑后 Confirm 保存。",
        "修改后务必「保存设置」或 Start+Select 写入 Flash，否则断电丢失。",
    ])

    # --- 8 Calibration ---
    add_title(doc, "八、校准指南", 1)

    add_title(doc, "8.1 屏幕瞄准校准（必做）", 2)
    add_p(doc, "用于将 IR 坐标映射到显示器区域，首次使用或换显示器后必须执行。")
    add_p(doc, "触发：首次开机按扳机；暂停菜单「屏幕校准」；Direct 暂停按扳机；OpenFIRE App 校准窗口。")
    add_p(doc, "OLED 校准界面顶部显示 校准: + 配置名。")
    add_p(doc, "步骤（每步对准屏幕指示点后再按扳机）：")
    add_bullets(doc, [
        "Top（顶部）→ Bottom（底部）→ Left（左侧）→ Right（右侧）",
        "Center（中心）— 保存四边 offset",
        "Verify（验证）— 试射满意后按扳机完成；不满意可按 A+B 重来",
        "取消：Reload+Home 恢复原校准数据",
    ])

    add_title(doc, "8.2 模拟摇杆校准", 2)
    add_table(doc,
              ["类型", "OLED 菜单/提示", "操作"],
              [
                  ["上电自动中心校准", "校准就绪", "每次开机约 2 秒，勿触碰摇杆"],
                  ["中心校准", "中心校准 → 中心已校", "将当前位置设为新中心"],
                  ["摇杆校准", "摇杆校准 → 摇杆:旋转/完成", "约 4 秒内画满圈后回中"],
              ])
    add_p(doc, "OpenFIRE App 目前无摇杆范围校准界面，请在光枪 OLED 暂停菜单「摇杆校准」中操作。")

    add_title(doc, "8.3 IR 发射器安装", 2)
    add_bullets(doc, [
        "方形布局：发射器安装在显示器上下（OLED 显示 方）。",
        "菱形布局：发射器安装在显示器四边中点（OLED 显示 菱）。",
        "OpenFIRE App：Open IR Emitter Alignment Assistant 可查看对齐示意图。",
    ])

    # --- 9 Dongle ---
    add_title(doc, "九、无线 Dongle 详细说明", 1)

    add_title(doc, "9.1 Dongle 是什么", 2)
    add_p(doc,
          "Dongle 是插在 PC 上的 ESP32-S3 接收器，通过 ESP-NOW 2.4GHz 与光枪通信。"
          "配对完成后向 PC 呈现与光枪相同的 USB 身份，并转发 HID 与串口力反馈数据。")

    add_title(doc, "9.2 屏幕状态对照（中文 OLED）", 2)
    add_table(doc,
              ["OLED 显示（中文）", "含义", "您应做什么"],
              [
                  ["OpenFIRE接收器", "启动 Logo", "等待"],
                  ["RF嗅探ch / WiFi扫描", "自动选择干扰最小信道", "勿拔 USB，约 10～15 秒"],
                  ["优选信道 / 请稍候", "信道评估中", "等待"],
                  ["信道xx配对 / 请开光枪 / ESP-NOW监听", "等待光枪配对", "打开光枪（勿插 PC USB）"],
                  ["射频已连接 / 等待数据流", "射频已连，等待 HID 流", "等待数秒"],
                  ["等待踏板", "等待无线踏板握手", "有踏板则打开；无踏板等待"],
                  ["玩家:N / 信道:N / 枪:MAC", "配对成功", "可在 PC 测试"],
              ])
    add_p(doc, "TFT 型 Dongle（LilyGO 等）可能显示：搜索中 / 优选信道 / WiFi扫描 / 连接中 / 等待踏板 / 玩家:N。")

    add_title(doc, "9.3 无线踏板", 2)
    add_bullets(doc, [
        "踏板本体无 OLED，配对时由光枪 OLED 顶部显示 踏台搜索xx秒。",
        "Dongle 配对光枪后可能显示 等待踏板，直至踏板无线握手完成。",
        "踏板仅提供 LED 状态，玩法与有线踏板相同。",
    ])

    add_title(doc, "9.4 PC 端 USB 与 COM 口", 2)
    add_bullets(doc, [
        "配对成功前 PC 可能看到占位设备；完成后变为光枪实际配置名称。",
        "力反馈（Mamehooker 等）：选择 Dongle 对应 COM 口，波特率 9600。",
        "改 VID/PID/模式请在光枪端设置，然后重启配对使 Dongle 更新。",
    ])

    add_title(doc, "9.5 多台设备建议", 2)
    add_bullets(doc, [
        "建议固定「1 枪 + 1 Dongle」；多人使用不同枪号与不同 Dongle。",
        "避免多个 Dongle 同时处于配对状态抢连同一光枪。",
        "更换 Dongle 后若仍连旧 MAC，需清除绑定或等待超时后全信道重搜。",
    ])

    # --- 10 App ---
    add_title(doc, "十、OpenFIRE-App-cn 配置软件（可选）", 1)
    add_bullets(doc, [
        "用途：引脚映射、开关选项、Profile 参数、USB P1～P4、IR 校准、按键测试、力反馈测试。",
        "连接：选择 COM 口，App 发送 Dock 命令进入 Docked 模式后同步数据。",
        "保存：修改后点击 Confirm。枪号在 Settings → USB Identity 选 P1～P4。",
        "固件需与 App 同属 OpenFIRE ESP32 中文分支（含 OLED_MENU_ZH），否则 App 可能拒绝连接。",
    ])

    # --- 11 MAME ---
    add_title(doc, "十一、力反馈与 MAMEHOOKER", 1)
    add_bullets(doc, [
        "编译启用 MAMEHOOKER 时，光枪/Dongle 提供 9600 波特率串口通道。",
        "PC 端软件发送力反馈指令 → 串口 →（无线时经 Dongle）→ 光枪驱动电磁阀/震动。",
        "OLED 可显示 Mamehook 弹药、生命等信息（取决于游戏配置）。",
        "暂停菜单「震动FFB开关」控制是否用震动电机代替电磁阀响应串口力反馈（与「震动开关」不同）。",
        "「电磁开关」控制电磁阀硬件；「震动开关」控制电机震动硬件。",
    ])

    # --- 12 Troubleshooting ---
    add_title(doc, "十二、常见问题与故障排除", 1)
    add_table(doc,
              ["现象", "可能原因", "处理办法"],
              [
                  ["Dongle 一直 请开光枪/搜索中", "光枪未开、仍插 USB、距离过远", "先 Dongle 后开枪；拔枪 USB；靠近 PC"],
                  ["光枪显示 踏台搜索 不停", "踏板未开机或距离过远", "打开踏板；靠近 Dongle"],
                  ["PC 设备名不变/仍是 P1", "枪号未保存或未重插 USB", "保存设置；重新插拔 USB"],
                  ["App 无法设置枪号", "固件与 App 协议不一致", "升级至最新中文 OLED 固件"],
                  ["无线连上后断线", "2.4G 干扰", "缩短距离；减少同频段设备"],
                  ["OLED 全黑", "I2C 地址/引脚与板型不符", "核对编译板型；重插 USB"],
                  ["瞄准偏移", "未校准或换显示器", "重做屏幕校准；检查方/菱与发射器安装"],
                  ["力反馈无效", "COM 口错误或 FFB 关闭", "选 Dongle COM；开启震动FFB开关；检查接线"],
                  ["游戏时间到 无法操作", "游戏计时到期锁定", "等待结束或将游戏计时设为 计时:关"],
                  ["上电摇杆漂移", "校准时碰了摇杆", "重启 2 秒内勿触碰；或做中心/摇杆校准"],
                  ["OLED 中文缺字", "字库未含该字符", "升级含最新字形的固件"],
              ])

    # --- 13 Appendix ---
    add_title(doc, "附录 A：推荐日常流程速查", 1)
    add_p(doc, "【有线游玩】USB 接 PC → 确认 USB 图标 → 进入游戏。", bold=True)
    add_p(doc, "【无线游玩】Dongle 接 PC → 光枪开机（不插 USB）→ 等 Wi-Fi 图标 → 进入游戏。", bold=True)
    add_p(doc, "【改配置/校准】Reload+Select 进暂停 → 操作 → 保存设置 或 Start+Select → Reload+Home 退出。", bold=True)
    add_p(doc, "【发 Esc】Reload+Start 或 Trigger+A+C。", bold=True)
    add_p(doc, "【多玩家】每只枪设不同枪号 → 保存设置 → 重插 USB / 各用 Dongle。", bold=True)

    add_title(doc, "附录 B：中文菜单与英文固件对照（简表）", 1)
    add_table(doc,
              ["中文（本固件）", "英文（原版）"],
              [
                  ["中心校准", "Center Calibrate"],
                  ["摇杆校准", "Range Calibrate"],
                  ["屏幕校准", "Calibrate"],
                  ["配置选择", "Profile Select"],
                  ["少键模式", "Low Button Toggle"],
                  ["震动FFB开关", "Rumble FFB Toggle"],
                  ["游戏计时", "Play Timer"],
                  ["玩具模式", "Boy Mode"],
                  ["枪号", "Gun ID"],
              ])

    add_title(doc, "附录 C：相关资源", 1)
    add_bullets(doc, [
        "固件仓库：OpenFIRE-Firmware-ESP32（分支 gamepad_new_20260606_zh）",
        "配置 App：OpenFIRE-App-cn",
        "板型引脚：docs/BOARDS_PIN.md",
        "原版项目：Team OpenFIRE",
    ])

    add_title(doc, "附录 D：版权声明", 1)
    add_p(doc,
          "OpenFIRE 光枪系统基于 Team OpenFIRE 开源项目。ESP32 移植、ESP-NOW 无线与中文 OLED 界面由 "
          "OpenFIRE-Firmware-ESP32 维护者实现。本说明书仅供用户参考，功能以实际刷入的固件版本与硬件配置为准。")

    doc.save(OUT)
    print("Saved:", OUT)


if __name__ == "__main__":
    build()
